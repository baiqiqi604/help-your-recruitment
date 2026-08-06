"""
简历写入模块

职责：
1. 将优化后的简历文本按小节写入结构化 docx（python-docx）
2. 支持 docx → PDF 转换（docx2pdf）

依赖：python-docx、docx2pdf
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

# 简历标准小节标题（按解析优先级排列）
SECTION_TITLES = [
    "姓名",
    "求职目标",
    "个人摘要",
    "核心技能",
    "工作经历",
    "项目经历",
    "教育背景",
    "证书",
]

# 常见标题前缀（如 "一、" "1." "-" 等），解析时剔除
_TITLE_PREFIX = re.compile(r"^[\s\d一二三四五六七八九十、.．\-—_*]*")


def _match_section_title(line: str) -> str | None:
    """判断一行是否为小节标题，返回标准标题；否则返回 None。"""
    cleaned = _TITLE_PREFIX.sub("", line.strip())
    for title in SECTION_TITLES:
        if cleaned.startswith(title):
            rest = cleaned[len(title):]
            # 标题行后面通常紧跟分隔符 / 冒号 / 竖线 / 换行
            if not rest or rest.startswith(("/", "：", ":", "、", "|", " ")):
                return title
            # 别名：如 "技能" → "核心技能"、"经历" → 忽略
    # 别名映射兜底
    aliases = {"技能": "核心技能", "联系方式": "姓名", "其他": "证书"}
    for alias, target in aliases.items():
        if cleaned.startswith(alias):
            return target
    return None


def _split_sections(text: str) -> list[tuple[str, list[str]]]:
    """将简历文本按小节标题拆分。

    Returns:
        形如 [("姓名", ["张三 | 138..."]), ("求职目标", [...])] 的列表
    """
    lines = [ln.strip() for ln in text.splitlines()]
    sections: list[tuple[str, list[str]]] = []
    current_title = "简历"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        if current_lines:
            sections.append((current_title, current_lines))
        current_lines = []

    for line in lines:
        if not line:
            continue
        title = _match_section_title(line)
        if title:
            flush()
            current_title = title
        else:
            current_lines.append(line)
    flush()

    if not sections:
        sections = [("简历", [ln for ln in lines if ln])]
    return sections


def write_optimized_resume(text: str, out_docx: str | Path) -> str:
    """将优化后的简历文本写入结构化 docx。

    按「姓名 / 求职目标 / 个人摘要 / 核心技能 / 工作经历 / 项目经历 /
    教育背景 / 证书」小节排版：小节标题加粗，首节（姓名）居中大号。

    Args:
        text: 优化后的简历全文
        out_docx: 输出 docx 路径

    Returns:
        输出 docx 的绝对路径

    Raises:
        ValueError: 简历文本为空
    """
    if not text or not text.strip():
        raise ValueError("简历文本不能为空")

    out = Path(out_docx)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)

    sections = _split_sections(text)

    for idx, (title, lines) in enumerate(sections):
        is_name_section = idx == 0 and title in ("姓名", "简历")

        # 小节标题
        heading = doc.add_paragraph()
        if is_name_section:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(16) if is_name_section else Pt(13)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(2)

        # 小节正文
        for line in lines:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            para.add_run(line)

    doc.save(str(out))
    logger.info("优化后简历已写入: %s（%d 个小节）", out, len(sections))
    return str(out)


def docx_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> str:
    """将 docx 文件转换为 PDF（docx2pdf，依赖本机 Word）。

    Args:
        docx_path: 源 docx 路径
        pdf_path: 输出 pdf 路径

    Returns:
        PDF 路径；转换失败（如未安装 MS Word）时记 warning 并返回 ""。
    """
    src = Path(docx_path)
    dst = Path(pdf_path)
    if not src.exists():
        raise ValueError(f"docx 文件不存在: {src}")

    try:
        from docx2pdf import convert

        dst.parent.mkdir(parents=True, exist_ok=True)
        convert(str(src), str(dst))
        logger.info("docx 转 PDF 完成: %s", dst)
        return str(dst)
    except Exception as e:  # noqa: BLE001
        logger.warning("docx 转 PDF 失败（可能需要安装 MS Word）: %s", e)
        return ""


if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO)
    sample = """姓名 / 联系方式
张三 | 13800000000 | zhangsan@example.com

求职目标
资深 Python 后端工程师

核心技能
Python、Django、MySQL、Redis

工作经历
某公司 后端工程师（2020-2023）
负责订单系统开发，日活提升 30%。"""
    out = sys.argv[1] if len(sys.argv) > 1 else "output/demo.docx"
    write_optimized_resume(sample, out)
