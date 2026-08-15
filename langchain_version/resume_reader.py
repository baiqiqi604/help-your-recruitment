"""
PDF 解析模块

职责：
1. 将 PDF 简历转换为 Word 文件（保留格式）
2. 从 Word 中提取段落文本和表格文本
3. 记录每个段落的样式信息（字体、字号、加粗等）

依赖：pdf2docx, python-docx
"""

from __future__ import annotations

import logging
import tempfile
import uuid
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """PDF 转 Word，保留原始排版。

    Args:
        pdf_path: 输入 PDF 文件路径
        docx_path: 输出 Word 文件路径

    Raises:
        FileNotFoundError: PDF 文件不存在
        ValueError: PDF 为空或无法解析（如扫描件、加密件）
    """
    pdf_file = Path(pdf_path)
    if not pdf_file.exists():
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    # 确保输出目录存在
    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)

    try:
        from pdf2docx import Converter
    except ImportError as e:
        raise ImportError("缺少依赖 pdf2docx，请执行: pip install pdf2docx") from e

    logger.info("开始转换 PDF -> Word: %s", pdf_path)
    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path)
    finally:
        cv.close()

    # 校验输出文件是否生成且非空
    out_file = Path(docx_path)
    if not out_file.exists() or out_file.stat().st_size == 0:
        raise ValueError(
            "PDF 转换失败或结果为空，可能是扫描件/加密件，暂不支持；"
            "请改传 DOCX 格式简历（复杂 PDF 无法高保真还原）"
        )
    logger.info("PDF 转换完成: %s", docx_path)


def read_resume(docx_path: str) -> dict[str, Any]:
    """读取 Word 简历内容。

    Args:
        docx_path: Word 文件路径

    Returns:
        {
            "paragraphs": [
                {"index": int, "text": str, "style": str, "font": {...}},
                ...
            ],
            "tables": [[["cell1", "cell2"], ...], ...],
            "full_text": str  # 纯文本拼接
        }

    Raises:
        FileNotFoundError: Word 文件不存在
    """
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")

    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    document = Document(docx_path)

    # 1. 读取段落（保留索引用于后续替换）
    paragraphs: list[dict[str, Any]] = []
    text_lines: list[str] = []
    for index, para in enumerate(document.paragraphs):
        text = para.text
        font_info = _extract_font_info(para)
        paragraphs.append(
            {
                "index": index,
                "text": text,
                "style": para.style.name if para.style else "",
                "font": font_info,
            }
        )
        if text.strip():
            text_lines.append(text)

    # 2. 读取表格（按行/列保留单元格结构）
    tables: list[list[list[str]]] = []
    for table in document.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
            # 表格内容也并入全文
            for cell_text in row_data:
                if cell_text.strip():
                    text_lines.append(cell_text)
        tables.append(table_data)

    # 3. 拼接全文
    full_text = "\n".join(text_lines)

    logger.info(
        "简历读取完成：%d 个段落，%d 个表格",
        len(paragraphs),
        len(tables),
    )

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
    }


def read_resume_text(path: str) -> str:
    """按扩展名读取简历纯文本（.txt 直读，.pdf 先转 docx 再提取，.docx 直接提取）。

    作为 main.py / web_app.py 的统一简历读取入口，消除两处重复的
    pdf_to_docx + read_resume 流程。

    Args:
        path: 简历文件路径（.pdf / .docx / .txt）

    Returns:
        简历全文纯文本

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的扩展名
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"简历文件不存在: {path}")

    suffix = file_path.suffix.lower()
    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return read_resume(str(file_path))["full_text"]
    if suffix == ".pdf":
        # 转 docx 写入系统临时目录（避免污染源文件所在目录），用完即删
        docx_path = Path(tempfile.gettempdir()) / f"{file_path.stem}_{uuid.uuid4().hex[:8]}.docx"
        pdf_to_docx(str(file_path), str(docx_path))
        try:
            return read_resume(str(docx_path))["full_text"]
        finally:
            docx_path.unlink(missing_ok=True)

    raise ValueError(f"不支持的简历文件类型: {suffix}（仅支持 .pdf / .docx / .txt）")


def _extract_font_info(paragraph) -> dict[str, Any]:
    """从段落的首个非空 run 提取字体信息。

    Returns:
        {"name": str, "size": float|None, "bold": bool, "italic": bool}
    """
    result: dict[str, Any] = {
        "name": None,
        "size": None,
        "bold": False,
        "italic": False,
    }
    # 取首个有文本的 run 作为格式代表
    target_run = None
    for run in paragraph.runs:
        if run.text.strip():
            target_run = run
            break
    if target_run is None and paragraph.runs:
        target_run = paragraph.runs[0]

    if target_run is not None:
        font = target_run.font
        result["name"] = font.name
        result["size"] = font.size.pt if font.size else None
        result["bold"] = bool(font.bold)
        result["italic"] = bool(font.italic)

    return result


if __name__ == "__main__":
    # 简单自测
    import sys

    if len(sys.argv) > 1:
        result = read_resume(sys.argv[1])
        print(result["full_text"][:500])
