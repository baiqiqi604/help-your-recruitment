"""
格式保留与输出模块

职责：
1. 将优化后的文本写回 Word 文件（保留原格式）
2. 可选：将 Word 导出为 PDF

替换策略：按行顺序替换 + 保留首 run 格式（最稳妥方案）

依赖：python-docx, docx2pdf
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def write_optimized_resume(
    original_docx: str, optimized_text: str, output_docx: str
) -> None:
    """把优化后的内容写回 Word，保留原有格式。

    策略：
    - 按段落顺序逐行替换文本
    - 保留每个段落首个 run 的字体格式
    - 行数不一致时：原简历行数多则截断，优化文本多则追加

    Args:
        original_docx: 原始 Word 文件路径（提供格式模板）
        optimized_text: 优化后的简历全文
        output_docx: 输出 Word 文件路径

    Raises:
        FileNotFoundError: 原始 Word 文件不存在
    """
    if not Path(original_docx).exists():
        raise FileNotFoundError(f"原始 Word 文件不存在: {original_docx}")

    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    # 确保输出目录存在
    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)

    document = Document(original_docx)

    # 将优化后文本按行拆分（过滤纯空行但保留顺序）
    new_lines = [line.rstrip() for line in optimized_text.splitlines()]
    # 去掉首尾空行
    while new_lines and not new_lines[0].strip():
        new_lines.pop(0)
    while new_lines and not new_lines[-1].strip():
        new_lines.pop()

    paragraphs = document.paragraphs
    logger.info(
        "写回简历：原段落 %d 个，优化后行数 %d",
        len(paragraphs),
        len(new_lines),
    )

    # 按行顺序替换：保留首 run 格式
    for i, para in enumerate(paragraphs):
        if i < len(new_lines):
            _replace_paragraph_text(para, new_lines[i])
        else:
            # 原简历段落多于优化文本：清空多余段落
            _replace_paragraph_text(para, "")

    # 优化文本行数多于原段落：在末尾追加新段落
    if len(new_lines) > len(paragraphs):
        for extra_line in new_lines[len(paragraphs):]:
            document.add_paragraph(extra_line)

    document.save(output_docx)
    logger.info("优化后简历已保存: %s", output_docx)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """替换段落文本，保留首个 run 的格式。

    策略：将新文本写入首个 run，清空其余 run 的内容。
    这样能保留首 run 的字体/字号/加粗等格式属性。
    """
    runs = paragraph.runs

    # 段落没有任何 run：直接新增一个
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return

    # 首 run 写入新文本
    runs[0].text = new_text
    # 清空其余 run（保留 run 对象但置空文本，避免破坏 XML 结构）
    for run in runs[1:]:
        run.text = ""


def docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """Word 转 PDF（可选功能）。

    转换策略：优先 docx2pdf（需本机 Microsoft Word，格式最保真）；
    未安装或转换失败时回退 pdf2docx 内置 Converter（纯 Python 实现，无需 Word）。

    Args:
        docx_path: 输入 Word 文件路径
        pdf_path: 输出 PDF 文件路径

    Raises:
        FileNotFoundError: 输入 Word 文件不存在
        ImportError: docx2pdf 与 pdf2docx 均不可用
        ValueError: 转换未产出 PDF 文件
    """
    if not Path(docx_path).exists():
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")

    # 确保输出目录存在
    Path(pdf_path).parent.mkdir(parents=True, exist_ok=True)

    # 优先 docx2pdf（依赖 MS Word，保真度高）
    try:
        from docx2pdf import convert

        logger.info("开始转换 Word -> PDF（docx2pdf）: %s", docx_path)
        convert(docx_path, pdf_path)
        if Path(pdf_path).exists():
            logger.info("PDF 导出完成: %s", pdf_path)
            return
        logger.warning("docx2pdf 未产出文件，回退 pdf2docx")
    except ImportError:
        logger.warning("docx2pdf 未安装，回退 pdf2docx")
    except Exception as e:  # noqa: BLE001
        logger.warning("docx2pdf 转换失败（%s），回退 pdf2docx", e)

    # 回退 pdf2docx（无需 Word）
    try:
        from pdf2docx import Converter
    except ImportError as e:
        raise ImportError(
            "缺少 PDF 导出依赖，请执行: pip install pdf2docx"
        ) from e

    logger.info("开始转换 Word -> PDF（pdf2docx）: %s", docx_path)
    converter = Converter(docx_path)
    try:
        converter.convert(pdf_path)
    finally:
        converter.close()

    if not Path(pdf_path).exists():
        raise ValueError("PDF 导出失败，docx2pdf 与 pdf2docx 均未产出文件")
    logger.info("PDF 导出完成: %s", pdf_path)


# ──────────────────────────────────────────────
# 定制化简历 / 面试建议 Word 文档生成（依据《定制化简历大师》Skill）
# ──────────────────────────────────────────────
def write_customized_resume(optimized_text: str, output_docx: str) -> str:
    """将优化后的简历文本生成为新的 Word 文档（不修改用户原始文件）。

    Args:
        optimized_text: 优化后的简历全文
        output_docx: 输出 docx 路径

    Returns:
        输出 docx 绝对路径

    Raises:
        ValueError: 简历文本为空
    """
    if not optimized_text or not optimized_text.strip():
        raise ValueError("简历文本不能为空")

    out = Path(output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)

    lines = [line.rstrip() for line in optimized_text.splitlines()]
    section_title = ("姓名", "联系方式", "求职目标", "个人摘要", "核心技能",
                     "工作经历", "项目经历", "教育背景", "证书", "奖项", "其他")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        para = doc.add_paragraph()
        run = para.add_run(stripped)
        if any(stripped.startswith(prefix) for prefix in section_title) and len(stripped) <= 30:
            run.bold = True
            run.font.size = Pt(13)
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(3)
        else:
            para.paragraph_format.space_after = Pt(2)

    doc.save(str(out))
    logger.info("定制化简历已保存: %s", out)
    return str(out)


def write_interview_advice_docx(advice_text: str, output_docx: str) -> str:
    """将面试建议（Markdown 文本）渲染为 Word 文档。

    支持 # / ## / ### 标题、- 无序列表、1. 有序列表、普通段落。

    Args:
        advice_text: 面试建议 Markdown 文本（来自 interview_advisor.build_interview_advice）
        output_docx: 输出 docx 路径

    Returns:
        输出 docx 绝对路径

    Raises:
        ValueError: 面试建议文本为空
    """
    if not advice_text or not advice_text.strip():
        raise ValueError("面试建议文本不能为空")

    out = Path(output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)

    for line in advice_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            para = doc.add_paragraph()
            run = para.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(3)
        elif stripped.startswith("## "):
            para = doc.add_paragraph()
            run = para.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("# "):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            para.paragraph_format.space_after = Pt(10)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.space_after = Pt(2)
            para.add_run("• " + stripped.lstrip("-• ").strip())
        elif stripped[0].isdigit() and ". " in stripped[:4]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.space_after = Pt(2)
            para.add_run(stripped)
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            para.add_run(stripped)

    doc.save(str(out))
    logger.info("面试建议已保存: %s", out)
    return str(out)


# ──────────────────────────────────────────────
# HTML 简历 / 结构化 YAML 输出（对应《resume-formatter》Skill）
# ──────────────────────────────────────────────
def write_customized_resume_html(
    optimized_text: str,
    output_html: str,
    output_yaml: str | None = None,
    template: str = "classic",
    output_docx: str | None = None,
) -> dict[str, str]:
    """将优化后的简历文本生成为精美 HTML 简历 + 结构化 YAML 数据文件。

    流程：纯文本 → LLM/规则结构化解析为 ResumeData → 按模板渲染为 HTML + YAML。
    解析失败不会抛异常，会退回最简规则结构，保证一定能产出文件。

    Args:
        optimized_text: 优化后的简历纯文本（来自 content_optimizer）
        output_html: 输出 HTML 文件路径（必填）
        output_yaml: 输出 YAML 数据文件路径（可选，默认在同目录下加 _data.yaml 后缀）
        template: 模板风格（classic / modern / professional / tech）
        output_docx: 可选，同时输出一页 A4 精美 Word 简历路径
            （与 HTML 共用同一次结构化解析，避免重复调用 LLM）

    Returns:
        {"html_path": 绝对路径, "yaml_path": 绝对路径, "check_report": 检查报告文本,
         "docx_path": 精美 Word 路径（传入 output_docx 时才有）}
    """
    from pathlib import Path as _Path

    from resume_formatter import (
        DEFAULT_TEMPLATE,
        format_check_report,
        fit_resume_to_one_page,
        parse_resume_text_to_data,
        render_resume_html,
        resume_to_yaml,
        run_resume_check,
    )

    if not optimized_text or not optimized_text.strip():
        raise ValueError("简历文本不能为空")

    # 解析为结构化数据
    data = parse_resume_text_to_data(optimized_text)
    tpl = template if template in {"classic", "modern", "professional", "tech"} else DEFAULT_TEMPLATE

    # 一页 A4 约束：渲染前对内容做温和裁剪（限制经历/项目段数与每段要点数），
    # 配合模板内嵌的自适应脚本（紧凑样式 + zoom 缩放），保证最终恰好一页 A4
    if tpl == "classic":
        data = fit_resume_to_one_page(data)

    # 写 HTML
    html_out = _Path(output_html)
    html_out.parent.mkdir(parents=True, exist_ok=True)
    html_out.write_text(render_resume_html(data, template=tpl), encoding="utf-8")
    logger.info("HTML 简历已保存: %s", html_out)

    # 写 YAML
    if output_yaml:
        yaml_out = _Path(output_yaml)
    else:
        yaml_out = html_out.with_name(html_out.stem + "_data.yaml")
    yaml_out.parent.mkdir(parents=True, exist_ok=True)
    yaml_out.write_text(resume_to_yaml(data), encoding="utf-8")
    logger.info("简历 YAML 数据已保存: %s", yaml_out)

    # 可选：同时输出一页 A4 精美 Word 简历（与 HTML 共用同一份裁剪后的 data）
    docx_path = ""
    if output_docx:
        try:
            from resume_writer import write_customized_resume_docx

            docx_path = write_customized_resume_docx(
                data, output_docx, fit_one_page=False
            )
        except Exception as e:  # noqa: BLE001
            logger.warning("精美 Word 简历生成失败（不影响 HTML/YAML）: %s", e)

    # 质量检查报告
    check_results = run_resume_check(data, resume_text=optimized_text)
    check_report = format_check_report(check_results)

    return {
        "html_path": str(html_out.resolve()),
        "yaml_path": str(yaml_out.resolve()),
        "check_report": check_report,
        "docx_path": docx_path,
    }


def write_customized_resume_docx(
    data: "ResumeData",
    output_docx: str,
    fit_one_page: bool = True,
) -> str:
    """基于结构化 ResumeData 生成精美 Word 简历（A4 一页约束）。

    版式对齐 classic HTML 模板：
    - A4 页面 + 紧凑页边距（上下 10/8mm，左右 14mm）
    - 姓名大号粗体 + 两列联系信息（无边框表格）
    - 灰色底纹章节标题条（#D8D9D8）
    - 教育/经历：左侧日期 + 右侧内容
    - 项目：粗体名称 + 右侧日期 + 要点
    - 微软雅黑 9pt 正文（较 HTML 10pt 更小，配合一页裁剪保证单页）

    Args:
        data: 结构化简历数据
        output_docx: 输出 docx 路径
        fit_one_page: 是否先执行一页 A4 内容裁剪（默认开启）

    Returns:
        输出 docx 绝对路径

    Raises:
        ValueError: data 为空
    """
    if data is None:
        raise ValueError("简历数据不能为空")

    out = Path(output_docx)
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    if fit_one_page:
        from resume_formatter import fit_resume_to_one_page

        data = fit_resume_to_one_page(data)

    b = data.basic
    name = b.name or "未命名"

    # ── 文档与页面设置：A4 + 紧凑边距（一页约束） ──
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(10)
    sec.bottom_margin = Mm(8)
    sec.left_margin = Mm(14)
    sec.right_margin = Mm(14)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.15

    def _set_font(run, size: int = 9, bold: bool = False, color: str | None = None) -> None:
        run.font.name = "微软雅黑"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _no_borders(table) -> None:
        tblPr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            borders.append(e)
        tblPr.append(borders)

    def _shade(paragraph, fill: str = "D8D9D8") -> None:
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def _add_section(title: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        _shade(p)
        _set_font(p.add_run(f"  {title}"), size=10.5, bold=True)

    def _add_cell_paragraph(cell, text: str, size: int = 9, bold: bool = False) -> None:
        p = cell.paragraphs[0] if not cell.paragraphs[0].runs else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _set_font(p.add_run(text), size=size, bold=bold)

    # ── 头部：姓名 + 两列联系信息 ──
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(4)
    _set_font(p_name.add_run(name), size=20, bold=True)

    header_left = []
    header_right = []
    if b.title:
        header_left.append(f"求职意向：{b.title}")
    if data.education and data.education[0].major:
        header_left.append(f"专业：{data.education[0].major}")
    if b.phone:
        header_left.append(f"电话：{b.phone}")
    if b.location:
        header_right.append(f"现居地址：{b.location}")
    if b.email:
        header_right.append(f"邮箱：{b.email}")
    if b.github:
        header_right.append(f"GitHub：{b.github}")

    if header_left or header_right:
        header_tbl = doc.add_table(rows=1, cols=2)
        header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _no_borders(header_tbl)
        left_cell, right_cell = header_tbl.rows[0].cells
        left_cell.width = Mm(88)
        right_cell.width = Mm(88)
        for i, line in enumerate(header_left):
            _add_cell_paragraph(left_cell, line)
        for i, line in enumerate(header_right):
            _add_cell_paragraph(right_cell, line)

    # ── 自我评价 ──
    if b.summary:
        _add_section("自我评价")
        import re as _re

        summary_lines = [s.strip() for s in b.summary.split("\n") if s.strip()]
        if len(summary_lines) <= 1:
            sentences = [s.strip() for s in _re.split(r"[。；！]", b.summary) if len(s.strip()) > 5]
            summary_lines = sentences if sentences else [b.summary]
        for line in summary_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            _set_font(p.add_run(f"• {line}"))

    # ── 教育背景 ──
    if data.education:
        _add_section("教育背景")
        for edu in data.education:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _no_borders(tbl)
            date_cell, content_cell = tbl.rows[0].cells
            date_cell.width = Mm(30)
            content_cell.width = Mm(146)
            if edu.period:
                _add_cell_paragraph(date_cell, edu.period)
            school_line = f"{edu.school}　　{edu.degree} {edu.major}".strip()
            _add_cell_paragraph(content_cell, school_line, bold=True)
            if edu.highlights:
                _add_cell_paragraph(content_cell, f"主修课程：{'、'.join(edu.highlights)}")
            if edu.gpa:
                _add_cell_paragraph(content_cell, f"GPA：{edu.gpa}")

    # ── 工作经历 / 学生工作经历 ──
    if data.experience:
        is_student_any = any(
            ("学生" in (e.company or "")) or ("校" in (e.company or "")) or ("学术" in (e.company or ""))
            for e in data.experience
        )
        _add_section("学生工作经历" if is_student_any else "工作经历")
        for exp in data.experience:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _no_borders(tbl)
            date_cell, content_cell = tbl.rows[0].cells
            date_cell.width = Mm(30)
            content_cell.width = Mm(146)
            if exp.period:
                _add_cell_paragraph(date_cell, exp.period)
            company_line = f"{exp.company}　　{exp.position}".strip()
            _add_cell_paragraph(content_cell, company_line, bold=True)
            for pt in exp.points:
                _add_cell_paragraph(content_cell, f"• {pt}")

    # ── 项目经历 ──
    if data.projects:
        _add_section("项目经历")
        for proj in data.projects:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _no_borders(tbl)
            date_cell, content_cell = tbl.rows[0].cells
            date_cell.width = Mm(30)
            content_cell.width = Mm(146)
            if proj.period:
                _add_cell_paragraph(date_cell, proj.period)
            proj_name = proj.name or ""
            role_tags = []
            if proj.role:
                role_tags.append(proj.role)
            if proj.tech_stack:
                role_tags.append("、".join(proj.tech_stack))
            title_text = proj_name + (f"　|　{'　'.join(role_tags)}" if role_tags else "")
            _add_cell_paragraph(content_cell, title_text, bold=True)
            for pt in proj.points:
                _add_cell_paragraph(content_cell, f"• {pt}")

    # ── 荣誉奖项 ──
    if data.awards:
        _add_section("荣誉奖项")
        for a in data.awards:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            text = a.name or ""
            if a.issuer:
                text += f"　　{a.issuer}"
            if a.date:
                text += f"　　{a.date}"
            _set_font(p.add_run(f"• {text}"))

    # ── 技能特长 ──
    if data.skills:
        _add_section("技能特长")
        for sc in data.skills:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            items = "、".join(sc.items or [])
            _set_font(p.add_run(f"• {sc.name}：{items}"))

    doc.save(str(out))
    logger.info("精美 Word 简历已保存: %s", out)
    return str(out)


if __name__ == "__main__":
    print("resume_writer 模块自测：需要传入实际文件路径")
