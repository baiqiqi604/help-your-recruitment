"""resume_reader / resume_writer 的确定性测试（python-docx 真实读写）。

用 python-docx 在临时目录生成真实 .docx 文件，验证：
  - 段落/表格/全文提取
  - 字体信息提取
  - 按行替换保留首 run 格式（含行数多/少两种路径）
  - 定制化简历与面试建议 Word 生成（标题/列表/编号渲染）
PDF 转换（pdf2docx / MS Word）依赖环境，仅测试错误路径与空输出校验。
"""

from __future__ import annotations

import sys
from typing import Any
from unittest import mock

import pytest

import resume_reader
import resume_writer


@pytest.fixture()
def docx_document():
    """创建一个带段落与表格的临时 docx，返回 (path, 期望文本片段)。"""
    from docx import Document

    doc = Document()
    p1 = doc.add_paragraph()
    run = p1.add_run("姓名：张三")
    run.bold = True
    doc.add_paragraph("工作经历：3 年 Python 后端")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "技能"
    table.cell(0, 1).text = "Python"
    table.cell(1, 0).text = "学历"
    table.cell(1, 1).text = "本科"
    return doc


# ──────────────────────────────────────────────
# resume_reader
# ──────────────────────────────────────────────
class TestResumeReader:
    def test_read_resume_paragraphs_tables_and_fulltext(self, docx_document, tmp_path) -> None:
        path = str(tmp_path / "resume.docx")
        docx_document.save(path)

        result = resume_reader.read_resume(path)
        assert len(result["paragraphs"]) == 2
        assert result["paragraphs"][0]["text"] == "姓名：张三"
        assert result["paragraphs"][0]["font"]["bold"] is True
        assert result["paragraphs"][0]["style"]
        assert len(result["tables"]) == 1
        assert result["tables"][0][0] == ["技能", "Python"]
        assert "姓名：张三" in result["full_text"]
        assert "Python" in result["full_text"]
        assert "本科" in result["full_text"]

    def test_read_resume_missing_file_raises(self) -> None:
        with pytest.raises(FileNotFoundError):
            resume_reader.read_resume("not_exist.docx")

    def test_extract_font_info_bold_run(self, docx_document) -> None:
        para = docx_document.paragraphs[0]
        info = resume_reader._extract_font_info(para)
        assert info["bold"] is True
        assert info["name"] is None or isinstance(info["name"], str)

    def test_extract_font_info_empty_paragraph(self) -> None:
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("")
        info = resume_reader._extract_font_info(para)
        assert info == {"name": None, "size": None, "bold": False, "italic": False}

    def test_pdf_to_docx_missing_file_raises(self, tmp_path) -> None:
        with pytest.raises(FileNotFoundError):
            resume_reader.pdf_to_docx("not_exist.pdf", str(tmp_path / "out.docx"))

    def test_pdf_to_docx_empty_output_raises(self, tmp_path, monkeypatch) -> None:
        """伪造 pdf2docx.Converter 不产出文件 → 触发 ValueError 校验。"""
        pdf_path = tmp_path / "fake.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")
        fake_module = mock.MagicMock()
        converter_cls = mock.MagicMock()
        converter_cls.return_value.convert.return_value = None
        fake_module.Converter = converter_cls
        monkeypatch.setitem(sys.modules, "pdf2docx", fake_module)

        with pytest.raises(ValueError, match="转换失败或结果为空"):
            resume_reader.pdf_to_docx(str(pdf_path), str(tmp_path / "out.docx"))
        converter_cls.return_value.close.assert_called_once()


# ──────────────────────────────────────────────
# resume_writer
# ──────────────────────────────────────────────
class TestWriteOptimizedResume:
    def _make_original(self, tmp_path) -> str:
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("姓名：张三")
        r.bold = True
        doc.add_paragraph("工作经历：订单系统开发")
        doc.add_paragraph("教育背景：本科")
        path = str(tmp_path / "original.docx")
        doc.save(path)
        return path

    def test_replace_lines_preserves_first_run_format(self, tmp_path) -> None:
        from docx import Document

        original = self._make_original(tmp_path)
        out = str(tmp_path / "out.docx")
        # 优化文本比原段落少一行 → 末尾段落被清空
        resume_writer.write_optimized_resume(original, "姓名：李四\n工作经历：高并发系统", out)

        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["姓名：李四", "工作经历：高并发系统", ""]
        # 首 run 格式保留（原段落 0 是加粗）
        assert doc.paragraphs[0].runs[0].bold is True

    def test_append_lines_when_optimized_longer(self, tmp_path) -> None:
        from docx import Document

        original = self._make_original(tmp_path)
        out = str(tmp_path / "out2.docx")
        resume_writer.write_optimized_resume(
            original, "姓名：李四\n经历1\n经历2\n经历3\n经历4", out
        )
        doc = Document(out)
        assert len(doc.paragraphs) == 5
        assert doc.paragraphs[-1].text == "经历4"

    def test_missing_original_raises(self, tmp_path) -> None:
        with pytest.raises(FileNotFoundError):
            resume_writer.write_optimized_resume("nope.docx", "文本", str(tmp_path / "x.docx"))

    def test_replace_paragraph_text_clears_extra_runs(self, docx_document) -> None:
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("旧文本A")
        para.add_run("旧文本B")
        resume_writer._replace_paragraph_text(para, "新文本")
        assert para.text == "新文本"
        assert para.runs[0].text == "新文本"
        assert para.runs[1].text == ""

    def test_replace_paragraph_text_empty_paragraph(self) -> None:
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        resume_writer._replace_paragraph_text(para, "内容")
        assert para.text == "内容"


class TestWriteCustomizedResume:
    def test_writes_docx_and_marks_section_titles(self, tmp_path) -> None:
        from docx import Document

        text = "姓名：张三\n联系方式：138xxxx\n个人摘要\n后端开发经验丰富"
        out = str(tmp_path / "customized.docx")
        result = resume_writer.write_customized_resume(text, out)
        assert result == out
        assert tmp_path.joinpath("customized.docx").exists()

        doc = Document(out)
        paragraphs = {p.text: p for p in doc.paragraphs}
        assert "姓名：张三" in paragraphs
        assert "后端开发经验丰富" in paragraphs
        # 节标题（"姓名" 开头且短）应加粗
        assert paragraphs["姓名：张三"].runs[0].bold is True
        # 普通行不加粗
        assert paragraphs["后端开发经验丰富"].runs[0].bold is not True

    def test_empty_text_raises(self, tmp_path) -> None:
        with pytest.raises(ValueError):
            resume_writer.write_customized_resume("   ", str(tmp_path / "x.docx"))


class TestWriteInterviewAdviceDocx:
    MARKDOWN = (
        "# 面试建议_某公司\n"
        "## 一、公司判断\n"
        "### 1. 概况\n"
        "- 业务增长稳健\n"
        "1. 建议准备 STAR 案例\n"
        "普通段落内容"
    )

    def test_renders_markdown_structure(self, tmp_path) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out = str(tmp_path / "advice.docx")
        resume_writer.write_interview_advice_docx(self.MARKDOWN, out)
        assert tmp_path.joinpath("advice.docx").exists()

        from docx import Document

        doc = Document(out)
        title = doc.paragraphs[0]
        assert title.text == "面试建议_某公司"
        assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert title.runs[0].bold is True
        assert title.runs[0].font.size.pt == 16

        h2 = doc.paragraphs[1]
        assert h2.text == "一、公司判断"
        assert h2.runs[0].bold is True
        assert h2.runs[0].font.size.pt == 14

        h3 = doc.paragraphs[2]
        assert h3.text == "1. 概况"
        assert h3.runs[0].bold is True

        bullet = doc.paragraphs[3]
        assert bullet.text == "• 业务增长稳健"

        numbered = doc.paragraphs[4]
        assert numbered.text == "1. 建议准备 STAR 案例"
        assert numbered.paragraph_format.left_indent is not None

        assert doc.paragraphs[5].text == "普通段落内容"

    def test_empty_text_raises(self, tmp_path) -> None:
        with pytest.raises(ValueError):
            resume_writer.write_interview_advice_docx("", str(tmp_path / "x.docx"))


class TestDocxToPdf:
    def test_missing_file_raises(self, tmp_path) -> None:
        with pytest.raises(FileNotFoundError):
            resume_writer.docx_to_pdf("nope.docx", str(tmp_path / "out.pdf"))
