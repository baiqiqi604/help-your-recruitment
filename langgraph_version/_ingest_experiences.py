"""将面试经验素材（JSON / txt / md / docx）结构化并入库 interview_kb。

用法:
    python _ingest_experiences.py <file> [<file> ...]   # 支持 .json / .txt / .md / .docx
    python _ingest_experiences.py                        # 缺省处理 data/raw/experience/ 下全部素材

说明:
    - .json 文件：按素材列表解析（experience_crawler 的存档格式）
    - .txt / .md / .docx 文件：整份文件作为一条手动面经素材（source=manual）入库
"""
from __future__ import annotations

import json
import logging
import sys
from datetime import datetime
from pathlib import Path

from experience_processor import process_raw_item
from interview_knowledge_base import add_experiences, count_questions

from config import PATH_CONFIG

logger = logging.getLogger(__name__)


def _read_docx_text(path: Path) -> str:
    """读取 Word 文档（.docx）的全文文本（段落 + 表格）。"""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("缺少依赖 python-docx，请执行: pip install python-docx") from e

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_items(file_path: str) -> list[dict]:
    """按文件扩展名加载素材列表。"""
    path = Path(file_path)
    if not path.exists():
        logger.warning("文件不存在: %s", file_path)
        return []
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ".docx"):
        # 整份文件作为一条手动面经素材
        if suffix == ".docx":
            text = _read_docx_text(path).strip()
        else:
            text = path.read_text(encoding="utf-8", errors="replace").strip()
        if not text:
            logger.warning("文件为空: %s", file_path)
            return []
        item = {
            "source": "manual",
            "title": f"手动导入_{path.stem}",
            "url": "",
            "content": text,
            "keyword": "",
            "company": "",
            "role": "",
            "stage": "",
            "collected_at": datetime.now().strftime("%Y-%m-%d"),
        }
        return [item]

    # 默认按 JSON 素材列表解析
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, list):
            return data
    except (OSError, json.JSONDecodeError) as e:
        logger.warning("解析失败 %s: %s", file_path, e)
    return []


def ingest_files(file_paths: list[str]) -> dict[str, int]:
    """将指定文件中的素材结构化入库，返回统计。"""
    items: list[dict] = []
    for fp in file_paths:
        items.extend(_load_items(fp))

    if not items:
        return {"素材": 0, "结构化": 0, "新增": 0, "题库总数": count_questions()}

    structured: list[dict] = []
    for item in items:
        try:
            structured.extend(process_raw_item(item))
        except Exception as e:  # noqa: BLE001
            logger.warning("结构化失败 %s: %s", item.get("title", ""), e)

    added = 0
    if structured:
        try:
            added = add_experiences(structured)
        except Exception as e:  # noqa: BLE001
            logger.warning("入库失败: %s", e)

    total = count_questions()
    logger.info("完成：素材 %d 条 → 结构化 %d 条 → 新增 %d 条（题库总数 %d）",
                len(items), len(structured), added, total)
    return {"素材": len(items), "结构化": len(structured), "新增": added, "题库总数": total}


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    args = sys.argv[1:]
    if args:
        result = ingest_files(args)
    else:
        exp_dir = Path(PATH_CONFIG["raw_data_dir"]) / "experience"
        files = sorted(exp_dir.glob("experiences_*.json")) + sorted(exp_dir.glob("manual_*.json"))
        result = ingest_files([str(f) for f in files])
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
